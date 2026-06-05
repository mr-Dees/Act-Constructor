"""
Тесты защиты DOCX-builder'а таблиц от объединений за границами (P6a: R6, DOCX-1).

- T6a.3: `_merge_cells` не должен ронять IndexError на объединении, выходящем
  за пределы матрицы (defense-in-depth поверх схемной валидации).
- T6a.4: бизнес-строка центрирования объединённой шапки вынесена в явный
  конфиг `CENTERED_MERGED_HEADER_TEXTS` в styles.py.
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.domains.acts.formatters.docx.builders.tables import build_table
from app.domains.acts.formatters.docx.styles import CENTERED_MERGED_HEADER_TEXTS
from app.domains.acts.schemas.act_content import TableCellSchema, TableSchema


def _ts(grid_data, **kw):
    """TableSchema из «сырого» списка строк (с обходом валидации схемы)."""
    grid = [[TableCellSchema(**c) for c in row] for row in grid_data]
    # model_construct минует валидаторы — нужно для теста out-of-bounds span,
    # который схема теперь отбраковала бы.
    return TableSchema.model_construct(id="t1", nodeId="n1", grid=grid, **kw)


# ── T6a.3: защита _merge_cells от выхода за границы ──


def test_merge_out_of_bounds_does_not_raise():
    """Объединение с colSpan/rowSpan за пределами матрицы не роняет builder.

    Инвариант guard'а (T6a.3): builder зажимает конец объединения в границы
    матрицы и НЕ падает IndexError'ом. Конкретный текст итоговой ячейки при
    схлопывании всей таблицы в одно объединение не специфицируется.
    """
    doc = Document()
    # colSpan=5 при 2 колонках, rowSpan=4 при 2 строках — out-of-bounds.
    schema = _ts([
        [{"content": "Шапка", "isHeader": True, "colSpan": 5, "rowSpan": 4},
         {"content": "B"}],
        [{"content": "1"}, {"content": "2"}],
    ])
    # Не должно бросить IndexError; документ собирается с таблицей 2×2.
    build_table(doc, schema)
    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 2


def test_normal_merge_still_works():
    """Корректное объединение (в пределах) по-прежнему работает после guard'а."""
    doc = Document()
    schema = _ts([
        [{"content": "Объединено", "isHeader": True, "colSpan": 2},
         {"content": "", "isSpanned": True}],
        [{"content": "1"}, {"content": "2"}],
    ])
    build_table(doc, schema)
    assert "Объединено" in doc.tables[0].rows[0].cells[0].text


# ── T6a.4: бизнес-строка вынесена в явный конфиг styles ──


def test_centered_merged_header_config_contains_known_string():
    assert "Количество клиентов / элементов, ед." in CENTERED_MERGED_HEADER_TEXTS


def test_merged_header_in_config_is_centered():
    """Объединённая шапка, чей текст в конфиге, выравнивается по центру."""
    doc = Document()
    schema = _ts([
        [{"content": "Количество клиентов / элементов, ед.",
          "isHeader": True, "colSpan": 2},
         {"content": "", "isSpanned": True}],
        [{"content": "a"}, {"content": "b"}],
    ])
    build_table(doc, schema)
    cell = doc.tables[0].rows[0].cells[0]
    assert cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_merged_header_not_in_config_is_justified():
    """Объединённая шапка, чей текст НЕ в конфиге, выравнивается по ширине."""
    doc = Document()
    schema = _ts([
        [{"content": "Прочие отклонения", "isHeader": True, "colSpan": 2},
         {"content": "", "isSpanned": True}],
        [{"content": "a"}, {"content": "b"}],
    ])
    build_table(doc, schema)
    cell = doc.tables[0].rows[0].cells[0]
    assert cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
