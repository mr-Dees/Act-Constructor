"""
Форматер для создания актов в формате DOCX.

Преобразует структуру акта в документ Microsoft Word с поддержкой
форматирования, таблиц, изображений и HTML-контента.
"""

import base64
import binascii
import gc
import logging
import re
import threading
import time
from html.parser import HTMLParser
from io import BytesIO

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.core.config import Settings
from app.formatters.base_formatter import BaseFormatter
from app.formatters.utils import HTMLUtils

logger = logging.getLogger("act_constructor.formatter")


class TimeoutError(Exception):
    """Исключение для операций, превысивших допустимое время выполнения."""
    pass


class InterruptibleParser:
    """
    Обертка для HTMLParser с возможностью прерывания по timeout.

    Позволяет безопасно остановить парсинг длинных или сложных
    HTML-документов, разбивая их на чанки.
    """

    def __init__(self, parser, timeout: int, chunk_size: int):
        """
        Инициализация прерываемого парсера.

        Args:
            parser: Экземпляр HTMLParser для обработки
            timeout: Максимальное время выполнения в секундах
            chunk_size: Размер чанка для парсинга в символах
        """
        self.parser = parser
        self.timeout = timeout
        self.chunk_size = chunk_size
        self.interrupted = False
        self.start_time = None

    def feed(self, data: str):
        """
        Парсит данные с проверкой timeout на каждом шаге.

        Args:
            data: HTML-строка для парсинга

        Raises:
            TimeoutError: Если парсинг превысил timeout
        """
        self.start_time = time.time()
        self.interrupted = False

        # Используем настраиваемый chunk_size
        for i in range(0, len(data), self.chunk_size):
            if self.interrupted:
                raise TimeoutError(f"Парсинг прерван по timeout {self.timeout}s")

            if time.time() - self.start_time > self.timeout:
                self.interrupted = True
                raise TimeoutError(f"Парсинг превысил timeout {self.timeout}s")

            chunk = data[i:i + self.chunk_size]
            self.parser.feed(chunk)

    def interrupt(self):
        """Прерывает текущий парсинг."""
        self.interrupted = True


class TimeoutContext:
    """
    Кроссплатформенный context manager с реальным прерыванием операций.

    Использует threading.Timer для ограничения времени выполнения
    блока кода. При превышении timeout прерывает операцию через
    переданный interruptible объект.
    """

    def __init__(self, seconds: int, interruptible=None):
        """
        Инициализация timeout контекста.

        Args:
            seconds: Максимальное время выполнения в секундах
            interruptible: Объект с методом interrupt() для прерывания
                операции
        """
        self.seconds = seconds
        self.timer = None
        self.timed_out = False
        self.interruptible = interruptible

    def _timeout_handler(self):
        """Обработчик срабатывания таймера с прерыванием операции."""
        self.timed_out = True
        if self.interruptible:
            self.interruptible.interrupt()
        logger.warning(f"Операция превысила timeout {self.seconds}s и была прервана")

    def __enter__(self):
        """Запускает таймер при входе в контекст."""
        self.timer = threading.Timer(self.seconds, self._timeout_handler)
        self.timer.daemon = True
        self.timer.start()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        """
        Останавливает таймер при выходе из контекста.

        Args:
            exc_type: Тип исключения (если было)
            exc_val: Значение исключения
            exc_tb: Traceback исключения

        Returns:
            False для проброса исключения

        Raises:
            TimeoutError: Если операция превысила timeout
        """
        if self.timer:
            self.timer.cancel()

        # Если был timeout, пробрасываем исключение
        if self.timed_out and exc_type is None:
            raise TimeoutError(f"Операция превысила timeout {self.seconds}s")

        return False


class HTMLToDocxParser(HTMLParser):
    """
    Парсер HTML для преобразования в Word runs.

    Поддерживает inline-форматирование (bold, italic, underline),
    гиперссылки, сноски и защиту от слишком глубокой вложенности.
    """

    def __init__(self, paragraph, document, max_depth: int = 100):
        """
        Инициализация парсера.

        Args:
            paragraph: Параграф DOCX для добавления runs
            document: Документ DOCX (для гиперссылок)
            max_depth: Максимальная глубина вложенности HTML
        """
        super().__init__()
        self.paragraph = paragraph
        self.document = document
        self.max_depth = max_depth
        self.current_depth = 0
        self._reset_state()

    def _reset_state(self):
        """Сбрасывает внутреннее состояние парсера."""
        self.bold = False
        self.italic = False
        self.underline = False
        self.strike = False
        self.font_size: int | None = None
        self.alignment: str | None = None
        self.text_buffer: list[str] = []
        self.in_div = False
        self.in_link = False
        self.link_url: str | None = None
        self.link_text: list[str] = []
        self.in_footnote = False
        self.footnote_text: str | None = None
        self.footnote_content: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str]]):
        """
        Обрабатывает открывающие HTML теги.

        Args:
            tag: Имя тега
            attrs: Список атрибутов тега

        Raises:
            ValueError: Если превышена максимальная глубина вложенности
        """
        # Проверка глубины вложенности для защиты от DoS
        self.current_depth += 1
        if self.current_depth > self.max_depth:
            logger.warning(f"Превышена максимальная глубина вложенности HTML: {self.max_depth}")
            raise ValueError(f"HTML слишком глубоко вложен (>{self.max_depth} уровней)")

        if tag == 'br':
            return

        self._flush_buffer()

        if tag in ['b', 'strong']:
            self.bold = True
        elif tag in ['i', 'em']:
            self.italic = True
        elif tag == 'u':
            self.underline = True
        elif tag in ['s', 'strike', 'del']:
            self.strike = True
        elif tag in ['span', 'div']:
            attrs_dict = dict(attrs)

            # Гиперссылки
            if 'data-link-url' in attrs_dict:
                self.in_link = True
                self.link_url = attrs_dict['data-link-url']
                self.link_text = []
                return

            # Сноски
            if 'data-footnote-text' in attrs_dict:
                self.in_footnote = True
                self.footnote_text = attrs_dict['data-footnote-text']
                self.footnote_content = []
                return

            # Парсим стили через утилиту
            if 'style' in attrs_dict:
                style_dict = HTMLUtils.parse_style_dict(attrs_dict['style'])

                if 'font-size' in style_dict:
                    size_str = style_dict['font-size'].replace('px', '').strip()
                    try:
                        self.font_size = int(size_str)
                    except ValueError:
                        pass

                if 'text-align' in style_dict:
                    self.alignment = style_dict['text-align']

            if tag == 'div':
                self.in_div = True

    def handle_endtag(self, tag: str):
        """
        Обрабатывает закрывающие HTML теги.

        Args:
            tag: Имя тега
        """
        # Уменьшаем счетчик глубины
        self.current_depth = max(0, self.current_depth - 1)

        if tag == 'br':
            self._flush_buffer()
            self.paragraph.add_run('\n')
            return

        self._flush_buffer()

        if tag in ['b', 'strong']:
            self.bold = False
        elif tag in ['i', 'em']:
            self.italic = False
        elif tag == 'u':
            self.underline = False
        elif tag in ['s', 'strike', 'del']:
            self.strike = False
        elif tag == 'span':
            if self.in_link:
                self._add_hyperlink()
                return
            if self.in_footnote:
                self._add_footnote()
                return
            self.font_size = None
        elif tag == 'div':
            self.paragraph.add_run('\n')
            self.in_div = False
            self.alignment = None

    def handle_data(self, data: str):
        """
        Обрабатывает текстовые данные между тегами.

        Args:
            data: Текстовое содержимое
        """
        if not data.strip():
            return

        if self.in_link:
            self.link_text.append(data)
            return

        if self.in_footnote:
            self.footnote_content.append(data)
            return

        self.text_buffer.append(data)

    def _flush_buffer(self):
        """Сбрасывает текстовый буфер в run с текущим форматированием."""
        if not self.text_buffer:
            return

        text = ''.join(self.text_buffer)
        self.text_buffer = []

        if not text.strip():
            return

        run = self.paragraph.add_run(text)
        self._apply_formatting(run)

    def _apply_formatting(self, run):
        """
        Применяет текущее форматирование к run.

        Args:
            run: Run объект DOCX
        """
        run.bold = self.bold
        run.italic = self.italic
        run.underline = self.underline

        if self.strike:
            run.font.strike = True

        if self.font_size:
            run.font.size = Pt(self.font_size)

    def _add_hyperlink(self):
        """Добавляет нативную гиперссылку Word с форматированием."""
        link_text = ''.join(self.link_text)

        try:
            part = self.paragraph.part
            r_id = part.relate_to(
                self.link_url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True
            )

            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)

            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')

            if self.bold:
                rPr.append(OxmlElement('w:b'))
            if self.italic:
                rPr.append(OxmlElement('w:i'))
            if self.underline:
                rPr.append(OxmlElement('w:u'))

            rStyle = OxmlElement('w:rStyle')
            rStyle.set(qn('w:val'), 'Hyperlink')
            rPr.append(rStyle)

            new_run.append(rPr)
            new_run.text = link_text
            hyperlink.append(new_run)

            self.paragraph._p.append(hyperlink)

            logger.debug(f"Добавлена гиперссылка: {link_text} -> {self.link_url}")
        except Exception as e:
            logger.warning(f"Ошибка создания гиперссылки: {e}")
            run = self.paragraph.add_run(f'{link_text} ({self.link_url})')
            self._apply_formatting(run)

        self.in_link = False
        self.link_url = None
        self.link_text = []

    def _add_footnote(self):
        """Добавляет сноску как надстрочный текст."""
        footnote_content = ''.join(self.footnote_content)

        run = self.paragraph.add_run(footnote_content)
        self._apply_formatting(run)

        footnote_run = self.paragraph.add_run(f' [{self.footnote_text}]')
        footnote_run.font.superscript = True
        footnote_run.font.size = Pt(8)

        logger.debug(f"Добавлена сноска: {footnote_content} -> {self.footnote_text}")

        self.in_footnote = False
        self.footnote_text = None
        self.footnote_content = []

    def close(self):
        """Завершает парсинг с полной очисткой состояния."""
        self._flush_buffer()
        super().close()

        # Явная очистка для предотвращения memory leak
        self._reset_state()
        self.current_depth = 0

    def reset(self):
        """Переопределяет reset для полной очистки состояния."""
        super().reset()
        self._reset_state()
        self.current_depth = 0


class DocxFormatter(BaseFormatter):
    """
    Форматер для преобразования акта в документ DOCX.

    Использует композицию HTMLUtils для работы с HTML и конфигурируемые
    параметры из Settings. Поддерживает таблицы, изображения, HTML-блоки
    и нарушения.
    """

    def __init__(self, settings: Settings):
        """
        Инициализация форматера с настройками.

        Args:
            settings: Глобальные настройки приложения
        """
        self.settings = settings

        # Загружаем константы из настроек
        self.MAX_HEADING_LEVEL = settings.docx_max_heading_level
        self.DEFAULT_IMAGE_WIDTH = Inches(settings.docx_image_width)
        self.CAPTION_FONT_SIZE = Pt(settings.docx_caption_font_size)
        self.MAX_IMAGE_SIZE_MB = settings.max_image_size_mb
        self.HTML_PARSE_TIMEOUT = settings.html_parse_timeout
        self.MAX_HTML_DEPTH = settings.max_html_depth
        self.HTML_PARSE_CHUNK_SIZE = settings.html_parse_chunk_size
        # Параметры retry логики
        self.MAX_RETRIES = settings.max_retries
        self.RETRY_DELAY = settings.retry_delay

        self.ALIGNMENT_MAP = {
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
            'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT
        }

        logger.debug(f"DocxFormatter инициализирован с настройками: "
                     f"image_width={settings.docx_image_width}\", "
                     f"max_image_size={settings.max_image_size_mb}MB, "
                     f"parse_timeout={settings.html_parse_timeout}s, "
                     f"chunk_size={settings.html_parse_chunk_size}, "
                     f"max_retries={settings.max_retries}")

    def format(self, data: dict) -> Document:
        """
        Форматирует данные акта в документ DOCX.

        Args:
            data: Данные акта (tree, tables, textBlocks, violations)

        Returns:
            Документ DOCX с отформатированным актом
        """
        logger.info("Начало форматирования акта в DOCX")
        doc = Document()

        violations = data.get('violations', {})
        textBlocks = data.get('textBlocks', {})
        tables = data.get('tables', {})

        heading = doc.add_heading('Акт', level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        tree = data.get('tree', {})
        root_children = tree.get('children', [])

        for item in root_children:
            self._add_item(doc, item, violations, textBlocks, tables, level=1)

        logger.info("Форматирование акта в DOCX завершено")
        return doc

    def _add_item(
            self,
            doc: Document,
            item: dict,
            violations: dict,
            textBlocks: dict,
            tables: dict,
            level: int = 1
    ):
        """
        Рекурсивно добавляет пункт акта в документ.

        Args:
            doc: Документ DOCX
            item: Узел дерева акта
            violations: Словарь нарушений
            textBlocks: Словарь текстовых блоков
            tables: Словарь таблиц
            level: Уровень вложенности (для заголовков)
        """
        label = item.get('label', '')
        item_type = item.get('type', 'item')

        if label and item_type not in ['textblock', 'violation']:
            heading_level = min(level, self.MAX_HEADING_LEVEL)
            doc.add_heading(label, level=heading_level)
            logger.debug(f"Добавлен заголовок уровня {heading_level}: {label}")

        content = item.get('content', '')
        if content:
            doc.add_paragraph(content)

        # Таблица
        table_id = item.get('tableId')
        if table_id and table_id in tables:
            try:
                self._add_table(doc, tables[table_id])
            except Exception as e:
                logger.exception(f"Ошибка добавления таблицы {table_id}: {e}")
                doc.add_paragraph(f"[Ошибка отображения таблицы: {table_id}]")

        # Текстовый блок
        textblock_id = item.get('textBlockId')
        if textblock_id and textblock_id in textBlocks:
            try:
                self._add_textblock(doc, textBlocks[textblock_id])
            except Exception as e:
                logger.exception(f"Ошибка добавления текстового блока {textblock_id}: {e}")
                doc.add_paragraph(f"[Ошибка отображения текстового блока: {textblock_id}]")

        # Нарушение
        violation_id = item.get('violationId')
        if violation_id and violation_id in violations:
            try:
                self._add_violation(doc, violations[violation_id])
            except Exception as e:
                logger.exception(f"Ошибка добавления нарушения {violation_id}: {e}")
                doc.add_paragraph(f"[Ошибка отображения нарушения: {violation_id}]")

        # Рекурсия для дочерних элементов
        children = item.get('children', [])
        for child in children:
            self._add_item(doc, child, violations, textBlocks, tables, level + 1)

    def _add_table(self, doc: Document, table_data: dict):
        """
        Добавляет таблицу в документ с поддержкой объединения ячеек.

        Args:
            doc: Документ DOCX
            table_data: Данные таблицы с grid структурой
        """
        grid = table_data.get('grid', [])

        if not grid or not grid[0]:
            doc.add_paragraph('[Пустая таблица]')
            logger.warning("Пропущена пустая таблица")
            return

        # Валидация матрицы
        num_cols = len(grid[0])
        for row_idx, row in enumerate(grid):
            if len(row) != num_cols:
                logger.error(f"Невалидная матрица таблицы: строка {row_idx}")
                doc.add_paragraph('[Ошибка структуры таблицы]')
                return

        num_rows = len(grid)
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'

        processed_merges = set()

        # Заполнение таблицы
        for row_idx, row_data in enumerate(grid):
            for col_idx, cell_data in enumerate(row_data):
                if cell_data.get('isSpanned', False):
                    continue

                cell = table.cell(row_idx, col_idx)
                cell.text = str(cell_data.get('content', ''))

                # Жирный заголовок
                if cell_data.get('isHeader', False):
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                # Объединение ячеек
                self._merge_cells(
                    table, cell, cell_data,
                    row_idx, col_idx,
                    num_rows, num_cols,
                    processed_merges
                )

        doc.add_paragraph()
        logger.debug(f"Добавлена таблица {num_rows}x{num_cols}")

    def _merge_cells(
            self, table, cell, cell_data: dict,
            row_idx: int, col_idx: int,
            num_rows: int, num_cols: int,
            processed_merges: set
    ):
        """
        Обрабатывает объединение ячеек (colspan/rowspan).

        Args:
            table: Таблица DOCX
            cell: Ячейка для объединения
            cell_data: Данные ячейки с rowSpan/colSpan
            row_idx: Индекс строки
            col_idx: Индекс столбца
            num_rows: Общее количество строк
            num_cols: Общее количество столбцов
            processed_merges: Set обработанных объединений
        """
        rowspan = cell_data.get('rowSpan', 1)
        colspan = cell_data.get('colSpan', 1)

        if rowspan > 1 or colspan > 1:
            merge_key = (row_idx, col_idx)
            if merge_key not in processed_merges:
                try:
                    end_row = min(row_idx + rowspan - 1, num_rows - 1)
                    end_col = min(col_idx + colspan - 1, num_cols - 1)
                    end_cell = table.cell(end_row, end_col)

                    cell.merge(end_cell)
                    processed_merges.add(merge_key)
                    logger.debug(f"Объединены ячейки [{row_idx},{col_idx}] -> [{end_row},{end_col}]")
                except Exception as e:
                    logger.exception(f"Ошибка объединения ячеек: {e}")

    def _add_textblock(self, doc: Document, textblock_data: dict):
        """
        Добавляет текстовый блок с HTML-форматированием и timeout защитой.

        Args:
            doc: Документ DOCX
            textblock_data: Данные блока с HTML контентом
        """
        content = textblock_data.get('content', '').strip()
        if not content:
            return

        formatting = textblock_data.get('formatting', {})
        blocks = self._parse_div_blocks(content, formatting)

        for block in blocks:
            if not block['content']:
                continue

            paragraph = doc.add_paragraph()
            paragraph.alignment = self.ALIGNMENT_MAP.get(
                block['alignment'],
                WD_PARAGRAPH_ALIGNMENT.LEFT
            )

            # Создаем новый парсер для каждого блока с защитой от
            # глубокой вложенности.
            parser = HTMLToDocxParser(paragraph, doc, max_depth=self.MAX_HTML_DEPTH)
            # Создаем прерываемый парсер с chunk_size из настроек
            interruptible = InterruptibleParser(
                parser,
                self.HTML_PARSE_TIMEOUT,
                self.HTML_PARSE_CHUNK_SIZE
            )

            try:
                # Применяем timeout с возможностью прерывания
                with TimeoutContext(self.HTML_PARSE_TIMEOUT, interruptible):
                    interruptible.feed(block['content'])
                    parser.close()
            except TimeoutError as e:
                logger.error(f"Timeout парсинга HTML: {e}")
                paragraph.add_run("[Контент слишком сложен для отображения]")
            except ValueError as e:
                # Слишком глубокая вложенность
                logger.error(f"HTML слишком вложен: {e}")
                paragraph.add_run("[Контент слишком сложен для отображения]")
            except Exception as e:
                logger.exception(f"Ошибка парсинга HTML: {e}")
                # Fallback: добавляем plain text
                clean_text = HTMLUtils.clean_html(block['content'])
                paragraph.add_run(clean_text)
            finally:
                # Гарантированная очистка памяти парсера
                try:
                    parser.reset()
                except Exception:
                    pass
                finally:
                    parser = None
                    interruptible = None
                    gc.collect()

            # Применяем базовый размер шрифта
            base_font_size = formatting.get('fontSize', 14)
            for run in paragraph.runs:
                if not run.font.size:
                    run.font.size = Pt(base_font_size)

        doc.add_paragraph()
        logger.debug("Добавлен текстовый блок")

    def _parse_div_blocks(self, content: str, formatting: dict) -> list[dict]:
        """
        Разбивает HTML на блоки по div-элементам с извлечением выравнивания.

        Args:
            content: HTML-контент
            formatting: Базовые параметры форматирования

        Returns:
            Список блоков с контентом и выравниванием
        """
        div_pattern = r'<div[^>]*>.*?</div>'
        div_matches = list(re.finditer(div_pattern, content, re.DOTALL))

        if not div_matches:
            return [{'content': content, 'alignment': formatting.get('alignment', 'left')}]

        blocks = []
        last_end = 0
        default_alignment = formatting.get('alignment', 'left')

        for match in div_matches:
            # Контент до div
            if match.start() > last_end:
                pre_content = content[last_end:match.start()].strip()
                if pre_content:
                    blocks.append({'content': pre_content, 'alignment': default_alignment})

            div_html = match.group(0)
            div_content = self._extract_div_content(div_html)

            # Используем утилиту HTMLUtils
            alignment = HTMLUtils.extract_style_property(
                div_html, 'text-align', default_alignment
            )

            if div_content:
                blocks.append({'content': div_content, 'alignment': alignment})

            last_end = match.end()

        # Контент после последнего div
        if last_end < len(content):
            post_content = content[last_end:].strip()
            if post_content:
                blocks.append({'content': post_content, 'alignment': default_alignment})

        return blocks

    def _extract_div_content(self, div_html: str) -> str:
        """
        Извлекает содержимое из div-тега.

        Args:
            div_html: HTML строка с div элементом

        Returns:
            Содержимое div без обрамляющих тегов
        """
        match = re.search(r'<div[^>]*>(.*?)</div>', div_html, re.DOTALL)
        return match.group(1).strip() if match else ''

    def _add_violation(self, doc: Document, violation_data: dict):
        """
        Добавляет блок нарушения со всеми секциями.

        Args:
            doc: Документ DOCX
            violation_data: Данные нарушения
        """
        self._add_labeled_field(doc, 'Нарушено', violation_data.get('violated', ''))
        self._add_labeled_field(doc, 'Установлено', violation_data.get('established', ''))
        self._add_description_list(doc, violation_data.get('descriptionList', {}))
        self._add_additional_content(doc, violation_data.get('additionalContent', {}))
        self._add_labeled_field(doc, 'Причины', violation_data.get('reasons', {}))
        self._add_labeled_field(doc, 'Последствия', violation_data.get('consequences', {}))
        self._add_labeled_field(doc, 'Ответственные', violation_data.get('responsible', {}))

        doc.add_paragraph()
        logger.debug("Добавлен блок нарушения")

    def _add_labeled_field(self, doc: Document, label: str, data):
        """
        Добавляет поле с жирной меткой.

        Args:
            doc: Документ DOCX
            label: Текст метки
            data: Данные поля (dict с enabled/content или строка)
        """
        if isinstance(data, dict):
            if not data.get('enabled', False):
                return
            content = data.get('content', '')
        else:
            content = data

        if content:
            p = doc.add_paragraph()
            p.add_run(f'{label}: ').bold = True
            p.add_run(content)

    def _add_description_list(self, doc: Document, desc_list: dict):
        """
        Добавляет список описаний в виде маркированного списка.

        Args:
            doc: Документ DOCX
            desc_list: Данные списка с items
        """
        if not desc_list.get('enabled', False):
            return

        items = desc_list.get('items', [])
        if not items:
            return

        p = doc.add_paragraph()
        p.add_run('Описание:').bold = True

        for item in items:
            if item.strip():
                doc.add_paragraph(item, style='List Bullet')

    def _add_additional_content(self, doc: Document, additional_content: dict):
        """
        Добавляет дополнительный контент (кейсы, изображения, свободный текст).

        Args:
            doc: Документ DOCX
            additional_content: Данные с items разных типов
        """
        if not additional_content.get('enabled', False):
            return

        items = additional_content.get('items', [])
        case_number = 1

        for item in items:
            item_type = item.get('type')

            if item_type == 'case':
                case_number = self._add_case(doc, item, case_number)
            elif item_type == 'image':
                self._add_image(doc, item)
                case_number = 1
            elif item_type == 'freeText':
                self._add_free_text(doc, item)
                case_number = 1

    def _add_case(self, doc: Document, item: dict, case_number: int) -> int:
        """
        Добавляет кейс с нумерацией.

        Args:
            doc: Документ DOCX
            item: Данные кейса
            case_number: Текущий номер кейса

        Returns:
            Следующий номер кейса
        """
        content = item.get('content', '')
        if content:
            p = doc.add_paragraph()
            p.add_run(f'Кейс {case_number}: ').bold = True
            p.add_run(content)
            return case_number + 1
        return case_number

    def _add_image(self, doc: Document, item: dict):
        """
        Добавляет изображение из base64 с управлением памятью, безопасной
        валидацией, retry логикой и graceful degradation.

        Args:
            doc: Документ DOCX
            item: Данные изображения (url, caption, filename)
        """
        url = item.get('url', '')
        caption = item.get('caption', '')
        filename = item.get('filename', '')

        if not (url and url.startswith('data:image')):
            logger.warning(f"Некорректный URL изображения: {filename}")
            self._add_image_fallback(doc, filename, caption, "некорректный URL")
            return

        image_stream = None
        image_data = None  # Отдельная переменная для контроля памяти
        retry_count = 0

        while retry_count <= self.MAX_RETRIES:
            try:
                # Безопасное разделение data URL
                if ',' not in url:
                    raise ValueError("Некорректный формат data URL")

                header, encoded = url.split(',', 1)

                # Безопасное декодирование base64 с обработкой ошибок
                try:
                    image_data = base64.b64decode(encoded, validate=True)
                except (binascii.Error, ValueError) as e:
                    logger.error(f"Ошибка декодирования base64 для {filename}: {e}")
                    self._add_image_fallback(doc, filename, caption, "некорректный base64")
                    return

                # Проверка размера
                size_mb = len(image_data) / (1024 * 1024)
                if size_mb > self.MAX_IMAGE_SIZE_MB:
                    logger.warning(f"Изображение {filename} слишком большое: {size_mb:.2f}MB")
                    self._add_image_fallback(doc, filename, caption, f"размер: {size_mb:.1f}MB")
                    return

                # Создаем stream внутри try-finally
                image_stream = BytesIO(image_data)

                # Вставка изображения
                img_paragraph = doc.add_paragraph()
                img_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                img_run = img_paragraph.add_run()
                img_run.add_picture(image_stream, width=self.DEFAULT_IMAGE_WIDTH)

                # Подпись
                if caption:
                    p = doc.add_paragraph(caption)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in p.runs:
                        run.italic = True
                        run.font.size = self.CAPTION_FONT_SIZE

                logger.debug(f"Добавлено изображение: {filename} ({size_mb:.2f}MB)")

                # Успешное добавление - выходим из retry loop
                break

            except (OSError, IOError) as e:
                # Временные ошибки ввода-вывода - пробуем retry
                retry_count += 1
                if retry_count <= self.MAX_RETRIES:
                    logger.warning(f"Ошибка I/O при добавлении изображения {filename} "
                                   f"(попытка {retry_count}/{self.MAX_RETRIES}): {e}")
                    time.sleep(self.RETRY_DELAY)
                    continue
                else:
                    logger.error(f"Исчерпаны попытки добавления изображения {filename}: {e}")
                    # Graceful degradation - показываем placeholder
                    self._add_image_placeholder(doc, filename, caption, "временная ошибка I/O")
                    break

            except Exception as e:
                # Неожиданные ошибки - graceful degradation без retry
                logger.exception(f"Неожиданная ошибка при добавлении изображения {filename}: {e}")
                self._add_image_placeholder(doc, filename, caption, str(e)[:100])
                break

            finally:
                # Явное освобождение памяти в любом случае
                if image_stream:
                    try:
                        image_stream.close()
                    except Exception:
                        pass
                    image_stream = None

                # Очищаем decoded data
                if image_data:
                    image_data = None

                # Принудительная сборка мусора
                gc.collect()

    def _add_image_fallback(self, doc: Document, filename: str, caption: str, reason: str = ""):
        """
        Добавляет текстовую ссылку при критической ошибке (не retry).

        Args:
            doc: Документ DOCX
            filename: Имя файла изображения
            caption: Подпись изображения
            reason: Причина ошибки
        """
        error_msg = f"Изображение: {filename}"
        if reason:
            error_msg += f" (ошибка: {reason})"
        p = doc.add_paragraph(error_msg)
        if caption:
            p.add_run(f" - {caption}")

    def _add_image_placeholder(self, doc: Document, filename: str, caption: str, reason: str = ""):
        """
        Добавляет placeholder изображения при временной ошибке.

        Graceful degradation: показываем рамку вместо изображения.

        Args:
            doc: Документ DOCX
            filename: Имя файла изображения
            caption: Подпись изображения
            reason: Причина недоступности
        """
        # Добавляем параграф с рамкой
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Текст placeholder
        placeholder_text = f"[Изображение: {filename}]"
        if reason:
            placeholder_text += f"\n(временно недоступно: {reason})"

        run = p.add_run(placeholder_text)
        run.font.size = Pt(10)
        run.italic = True

        # Добавляем подпись если есть
        if caption:
            caption_p = doc.add_paragraph(caption)
            caption_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in caption_p.runs:
                run.italic = True
                run.font.size = self.CAPTION_FONT_SIZE

        logger.info(f"Добавлен placeholder для изображения: {filename}")

    def _add_free_text(self, doc: Document, item: dict):
        """
        Добавляет свободный текст.

        Args:
            doc: Документ DOCX
            item: Данные с текстом
        """
        content = item.get('content', '')
        if content:
            doc.add_paragraph(content)
